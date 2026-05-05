"""
Documentos de treino dos agentes — extração, chunking e indexação no ChromaDB.

Embeddings rodam *no backend* (sentence-transformers/all-MiniLM-L6-v2) e os
vetores prontos são enviados ao Chroma. Isso mantém o container do Chroma
enxuto e nos dá controle total sobre o modelo de embedding.

Layout:
- arquivos físicos:  /data/uploads/agents/{agent_id}/{document_id}__{filename}
- coleção no Chroma: agent_{agent_id}
- chunk id no Chroma: {document_id}::{chunk_index}
"""
from __future__ import annotations
import logging
import os
from pathlib import Path
from typing import Iterable, Optional

from sqlalchemy.orm import Session

from models.agent import Agent, AgentDocument

log = logging.getLogger(__name__)

# ───────────────────────── config ─────────────────────────

UPLOADS_ROOT       = Path(os.getenv("UPLOADS_PATH", "/data/uploads/agents"))
CHROMA_HOST        = os.getenv("CHROMA_HOST", "chroma")
CHROMA_PORT        = int(os.getenv("CHROMA_PORT", "8000"))
EMBEDDING_MODEL    = os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
CHUNK_TOKENS       = 500
CHUNK_OVERLAP      = 50
TOKEN_ENCODING     = "cl100k_base"  # estimador suficiente para chunking

ALLOWED_MIME = {
    "application/pdf":                                                          "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document":  "docx",
    "application/msword":                                                       "docx",
    "text/plain":                                                               "txt",
    "text/markdown":                                                            "txt",
}


# ───────────────────── lazy-loaded singletons ─────────────────────
# Bibliotecas pesadas (sentence-transformers, chromadb) só carregam na 1ª chamada
# real — assim os testes podem fazer mock antes do load e a app sobe rápido.

_chroma_client = None
_embedder      = None
_encoder       = None


def _get_chroma():
    global _chroma_client
    if _chroma_client is None:
        import chromadb
        _chroma_client = chromadb.HttpClient(host=CHROMA_HOST, port=CHROMA_PORT)
    return _chroma_client


def _get_embedder():
    global _embedder
    if _embedder is None:
        from sentence_transformers import SentenceTransformer
        log.info("Carregando modelo de embeddings %s …", EMBEDDING_MODEL)
        _embedder = SentenceTransformer(EMBEDDING_MODEL)
    return _embedder


def _get_token_encoder():
    global _encoder
    if _encoder is None:
        import tiktoken
        _encoder = tiktoken.get_encoding(TOKEN_ENCODING)
    return _encoder


def _collection_name(agent_id: str) -> str:
    return f"agent_{agent_id}"


def _get_or_create_collection(agent_id: str):
    return _get_chroma().get_or_create_collection(
        name=_collection_name(agent_id),
        metadata={"hnsw:space": "cosine"},
    )


# ───────────────────────── extração de texto ─────────────────────────

def extract_text(file_path: str | Path, mime_type: Optional[str]) -> str:
    file_path = Path(file_path)
    kind = ALLOWED_MIME.get(mime_type or "", None)
    if kind is None:
        # fallback por extensão
        ext = file_path.suffix.lower().lstrip(".")
        kind = {"pdf": "pdf", "docx": "docx", "doc": "docx", "txt": "txt", "md": "txt"}.get(ext)

    if kind == "pdf":
        import pdfplumber
        parts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                if txt.strip():
                    parts.append(txt)
        return "\n\n".join(parts)

    if kind == "docx":
        import docx  # python-docx
        d = docx.Document(str(file_path))
        return "\n".join(p.text for p in d.paragraphs if p.text)

    if kind == "txt":
        return file_path.read_text(encoding="utf-8", errors="replace")

    raise ValueError(f"Tipo de arquivo não suportado: mime={mime_type} ext={file_path.suffix}")


# ───────────────────────── chunking ─────────────────────────

def chunk_text(text: str, chunk_tokens: int = CHUNK_TOKENS, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Slice por tokens com overlap. Tokens via tiktoken cl100k_base (estimativa OK para PT-BR)."""
    if not text or not text.strip():
        return []
    enc = _get_token_encoder()
    tokens = enc.encode(text)
    if not tokens:
        return []

    step = max(1, chunk_tokens - overlap)
    chunks: list[str] = []
    for start in range(0, len(tokens), step):
        end = start + chunk_tokens
        piece = enc.decode(tokens[start:end])
        if piece.strip():
            chunks.append(piece)
        if end >= len(tokens):
            break
    return chunks


def count_tokens(text: str) -> int:
    return len(_get_token_encoder().encode(text or ""))


# ───────────────────────── indexação ─────────────────────────

def index_to_chroma(agent_id: str, document_id: str, chunks: list[str], filename: str) -> None:
    """Calcula embeddings localmente e faz upsert na collection do agente."""
    if not chunks:
        return
    embedder = _get_embedder()
    embeddings = embedder.encode(chunks, convert_to_numpy=True, show_progress_bar=False).tolist()
    ids        = [f"{document_id}::{i}" for i in range(len(chunks))]
    metadatas  = [{"document_id": document_id, "filename": filename, "chunk_index": i}
                  for i in range(len(chunks))]

    coll = _get_or_create_collection(agent_id)
    coll.upsert(ids=ids, documents=chunks, embeddings=embeddings, metadatas=metadatas)


def _delete_chunks_from_chroma(agent_id: str, document_id: str) -> None:
    try:
        coll = _get_or_create_collection(agent_id)
        coll.delete(where={"document_id": document_id})
    except Exception as e:
        log.warning("Falha ao remover chunks do Chroma (agent=%s doc=%s): %s",
                    agent_id, document_id, e)


def drop_agent_collection(agent_id: str) -> None:
    """Remove a collection inteira — usado quando o agente é apagado."""
    try:
        _get_chroma().delete_collection(_collection_name(agent_id))
    except Exception as e:
        log.warning("Falha ao apagar collection agent_%s: %s", agent_id, e)


# ───────────────────────── upload + processamento ─────────────────────────

def _agent_upload_dir(agent_id: str) -> Path:
    p = UPLOADS_ROOT / agent_id
    p.mkdir(parents=True, exist_ok=True)
    return p


def upload_document(
    db: Session,
    agent_id: str,
    filename: str,
    content: bytes,
    mime_type: Optional[str],
) -> Optional[AgentDocument]:
    """Salva o arquivo, registra no banco com status=processing, processa e indexa."""
    agent = db.get(Agent, agent_id)
    if not agent:
        return None

    doc = AgentDocument(
        agent_id=agent_id,
        filename=filename,
        file_path="",  # preenchido após gerar id
        mime_type=mime_type,
        status="processing",
    )
    db.add(doc)
    db.flush()  # gera id

    safe_name = filename.replace("/", "_").replace("\\", "_")
    target = _agent_upload_dir(agent_id) / f"{doc.id}__{safe_name}"
    target.write_bytes(content)
    doc.file_path = str(target)
    db.commit()
    db.refresh(doc)

    try:
        text = extract_text(target, mime_type)
        chunks = chunk_text(text)
        index_to_chroma(agent_id, doc.id, chunks, filename)

        doc.chunks_count = len(chunks)
        doc.total_tokens = count_tokens(text)
        doc.status = "ready"
        doc.error_message = None
    except Exception as e:
        log.exception("Falha ao processar documento %s do agente %s", doc.id, agent_id)
        doc.status = "failed"
        doc.error_message = str(e)[:500]

    db.commit()
    db.refresh(doc)
    return doc


def list_documents(db: Session, agent_id: str) -> list[AgentDocument]:
    return (
        db.query(AgentDocument)
        .filter(AgentDocument.agent_id == agent_id)
        .order_by(AgentDocument.created_at.desc())
        .all()
    )


def get_document(db: Session, document_id: str) -> Optional[AgentDocument]:
    return db.get(AgentDocument, document_id)


def delete_document(db: Session, document_id: str) -> bool:
    doc = db.get(AgentDocument, document_id)
    if not doc:
        return False

    # Remove vetores no Chroma
    _delete_chunks_from_chroma(doc.agent_id, doc.id)

    # Remove arquivo físico
    if doc.file_path:
        try:
            Path(doc.file_path).unlink(missing_ok=True)
        except Exception as e:
            log.warning("Falha ao remover arquivo %s: %s", doc.file_path, e)

    db.delete(doc)
    db.commit()
    return True


__all__: Iterable[str] = (
    "upload_document",
    "list_documents",
    "get_document",
    "delete_document",
    "extract_text",
    "chunk_text",
    "count_tokens",
    "index_to_chroma",
    "drop_agent_collection",
)
